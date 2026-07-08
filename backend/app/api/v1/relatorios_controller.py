from flask import Blueprint, jsonify, request
from flask_jwt_extended import jwt_required
from app.core.database import db
from app.models.pedido import Pedido
from app.models.evento import Evento
from sqlalchemy import func
import io

relatorios_bp = Blueprint('relatorios', __name__)

@relatorios_bp.route('/dashboard', methods=['GET'])
@jwt_required()
def get_dashboard():
    vendas = db.session.query(func.count(Pedido.id)).scalar() or 0
    eventos = db.session.query(func.count(Evento.id)).scalar() or 0
    receita_total = db.session.query(func.sum(Pedido.valor_total)).scalar() or 0
    
    # Example: group by real month, assuming SQLite or MySQL works. In MySQL: func.month(Pedido.data_pedido).
    # Since we want it compatible, we can just fetch all and group in Python for safety, or use func.strftime('%m', Pedido.data_pedido) in SQLite, 
    # but the instructions say we use real MySQL. In MySQL it's func.month() or extract('month', ...).
    # To keep it simple and robust across backends without dialects, let's fetch in last 6 months.
    from datetime import datetime, timedelta
    
    six_months_ago = datetime.utcnow() - timedelta(days=180)
    recent_pedidos = db.session.query(Pedido.data_pedido, Pedido.valor_total).filter(Pedido.data_pedido >= six_months_ago).all()
    
    monthly_sales = {}
    for p in recent_pedidos:
        mes_str = p.data_pedido.strftime("%Y-%m")
        monthly_sales[mes_str] = monthly_sales.get(mes_str, 0) + float(p.valor_total or 0)
        
    grafico = [{"mes": m, "valor": v} for m, v in sorted(monthly_sales.items())]
    
    # Extra KPIs
    from app.models.pedido import EstadoPedido
    from app.models.ordem_producao import OrdemProducao, EstadoProducao
    from app.models.ocorrencia import OcorrenciaMaterial, TipoOcorrencia
    
    pedidos_pendentes = db.session.query(func.count(Pedido.id)).filter(Pedido.estado.in_([EstadoPedido.AGENDADO.value, EstadoPedido.CONFIRMADO.value])).scalar() or 0
    ordens_ativas = db.session.query(func.count(OrdemProducao.id)).filter(OrdemProducao.estado.in_([EstadoProducao.PENDENTE.value, EstadoProducao.EM_PRODUCAO.value])).scalar() or 0
    
    return jsonify({
        "success": True,
        "data": {
            "kpis": {
                "total_vendas": vendas,
                "total_eventos": eventos,
                "receita_estimada": float(receita_total),
                "pedidos_pendentes": pedidos_pendentes,
                "ordens_ativas": ordens_ativas
            },
            "graficos": {
                "vendas_por_mes": grafico
            }
        }
    }), 200

@relatorios_bp.route('/exportar/<tipo>', methods=['GET'])
@jwt_required()
def exportar(tipo):
    format_exc = request.args.get('formato', 'pdf').lower() # pdf, word, excel, csv
    
    from flask import send_file
    buffer = io.BytesIO()
    
    titulo = f"Relatório de {tipo.upper()}"
    ext = format_exc
    mimetype = 'text/plain'
    
    # Simple data generator for demo purposes
    data_mock = [
        ["ID", "Data", "Valor", "Status"],
        [1, "2023-01-01", "100.00", "Concluído"],
        [2, "2023-01-02", "150.00", "Pendente"]
    ]
    
    if format_exc == 'pdf':
        try:
            from reportlab.pdfgen import canvas
            from reportlab.lib.pagesizes import letter
            c = canvas.Canvas(buffer, pagesize=letter)
            c.drawString(100, 750, titulo)
            y = 700
            for row in data_mock:
                c.drawString(100, y, " | ".join(str(item) for item in row))
                y -= 20
            c.showPage()
            c.save()
            mimetype = 'application/pdf'
        except ImportError:
            buffer.write(b"reportlab nao instalado")
    
    elif format_exc == 'excel':
        try:
            from openpyxl import Workbook
            wb = Workbook()
            ws = wb.active
            ws.title = "Relatorio"
            for row in data_mock:
                ws.append(row)
            wb.save(buffer)
            mimetype = 'application/vnd.openxmlformats-officedocument.spreadsheetml.sheet'
            ext = 'xlsx'
        except ImportError:
            buffer.write(b"openpyxl nao instalado")
    
    elif format_exc == 'word':
        try:
            import docx
            doc = docx.Document()
            doc.add_heading(titulo, 0)
            table = doc.add_table(rows=1, cols=len(data_mock[0]))
            hdr_cells = table.rows[0].cells
            for i, val in enumerate(data_mock[0]):
                hdr_cells[i].text = str(val)
            for row_data in data_mock[1:]:
                row_cells = table.add_row().cells
                for i, val in enumerate(row_data):
                    row_cells[i].text = str(val)
            doc.save(buffer)
            mimetype = 'application/vnd.openxmlformats-officedocument.wordprocessingml.document'
            ext = 'docx'
        except ImportError:
            buffer.write(b"python-docx nao instalado")
            
    elif format_exc == 'csv':
        import csv
        # For CSV we need to handle strings not bytes first, then encode
        output = io.StringIO()
        writer = csv.writer(output)
        writer.writerows(data_mock)
        buffer.write(output.getvalue().encode('utf-8'))
        mimetype = 'text/csv'
        
    else:
        buffer.write(f"Relatório de {tipo.upper()} em formato {format_exc.upper()}".encode('utf-8'))
        
    buffer.seek(0)
    
    return send_file(
        buffer, 
        as_attachment=True, 
        download_name=f"relatorio_{tipo}.{ext}", 
        mimetype=mimetype
    )
